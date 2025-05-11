import os
import re
import logging
import pandas as pd
from datetime import datetime
from docx import Document
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import PyPDF2
import configparser
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from openai import OpenAI
import json
import time
import subprocess
from typing import Dict, List, Optional, Tuple, Any
import shutil
from typing import List, Dict
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

# 配置日志
logging.basicConfig(
    filename='recruitment_system.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filemode='a'
)

class ConfigManager:
    """配置管理器"""
    def __init__(self, config_file='config.ini'):
        self.config_file = config_file
        self.config = configparser.ConfigParser()
        self._load_config()

    def _load_config(self):
        if not os.path.exists(self.config_file):
            self._create_default_config()
        self.config.read(self.config_file, encoding='utf-8')

    def _create_default_config(self):
        self.config['DEFAULT'] = {
            'api_key': 'your_api_key_here',
            'base_url': 'https://api.deepseek.com',
            'tesseract_path': r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            'poppler_path': r'C:\Users\53115\AppData\Local\poppler-24.08.0\Library\bin'
        }
        self.config['PATHS'] = {
            'work_dir': '',
            'resume_dir': '',
            'job_desc_dir': '',
            'output_excel': '简历信息一览表.xlsx'
        }
        with open(self.config_file, 'w', encoding='utf-8') as f:
            self.config.write(f)

    def get(self, section: str, key: str) -> str:
        try:
            value = self.config[section][key]
            # 验证output_excel是否有效
            if key == 'output_excel' and value and not value.lower().endswith(('.xlsx', '.xlsm', '.xltx', '.xltm')):
                logging.warning(f"无效的Excel文件路径在config.ini: {value}")
                return self.config['PATHS']['output_excel']  # 回退到默认值
            return value
        except KeyError:
            return self.config['DEFAULT'].get(key, '')

    def set(self, section: str, key: str, value: str):
        if section not in self.config:
            self.config[section] = {}
        self.config[section][key] = value
        with open(self.config_file, 'w', encoding='utf-8') as f:
            self.config.write(f)

class JobDescriptionProcessor:
    """职位说明书处理器：负责提取岗位名称和完整内容"""
    
    def __init__(self):
        self.config = ConfigManager()
        self.client = OpenAI(
            api_key=self.config.get('API', 'api_key'),
            base_url=self.config.get('API', 'base_url')
        )
        self.job_cache = {}  # 缓存职位说明书信息 {文件名: {position, content}}
        self.resume_processor = ResumeProcessor()

    def extract_job_position_and_content(self, jd_text: str, jd_filename: str) -> Tuple[str, str]:
        """从职位说明书中提取岗位名称和完整内容"""
        prompt = """请从以下职位说明书中提取招聘岗位名称和完整内容。
要求：
1. 从"职位名称"、"岗位名称"、"招聘岗位"等字段中提取岗位名称。
2. 返回完整的职位说明书内容（包括任职资格、教育背景、技能要求等）。
3. 如果无法提取岗位名称，从文件名推断。
4. 返回格式：
   岗位名称：[提取的岗位名称]
   完整内容：[完整的职位说明书内容]

职位说明书内容：
{j}

示例返回：
岗位名称：软件工程师
完整内容：岗位名称：软件工程师\n任职资格：本科及以上，3年开发经验，熟悉Python...
""".format(j=jd_text)

        try:
            response = self.client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "你是一个专业的职位说明书分析专家。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1
            )

            content = response.choices[0].message.content.strip()
            if content:
                position, full_content = self._parse_position_and_content(content, jd_filename)
                logging.info(f"提取岗位名称: {position}, 内容长度: {len(full_content)} ({jd_filename})")
                return position, full_content
            else:
                logging.warning(f"未提取到岗位名称或内容: {jd_filename}")
                position = re.sub(r'职位说明书$|岗位说明书$', '', jd_filename).strip()
                return position, jd_text

        except Exception as e:
            logging.error(f"提取职位说明书信息失败: {jd_filename} - {str(e)}")
            position = re.sub(r'职位说明书$|岗位说明书$', '', jd_filename).strip()
            return position, jd_text

    def _parse_position_and_content(self, content: str, jd_filename: str) -> Tuple[str, str]:
        """解析岗位名称和完整内容"""
        parts = content.split('完整内容：', 1)
        if len(parts) == 2:
            position = parts[0].replace("岗位名称：", "").strip()
            full_content = parts[1].strip()
            return position, full_content
        logging.warning(f"解析岗位名称失败，使用文件名: {jd_filename}")
        position = re.sub(r'职位说明书$|岗位说明书$', '', jd_filename).strip()
        return position, content

    def process_job_descriptions(self, job_desc_files: List[str]) -> Dict[str, Dict[str, str]]:
        """处理所有职位说明书并缓存"""
        self.job_cache.clear()

        for jd_file in job_desc_files:
            try:
                jd_text = self.resume_processor.extract_text_from_docx(jd_file)
                if not jd_text.strip():
                    logging.warning(f"职位说明书内容为空: {jd_file}")
                    continue

                position, full_content = self.extract_job_position_and_content(jd_text, os.path.basename(jd_file))
                if position:
                    self.job_cache[jd_file] = {
                        'position': position,
                        'content': full_content
                    }
                    logging.info(f"成功缓存职位说明书: {jd_file}")
                else:
                    logging.warning(f"跳过无有效岗位名称的职位说明书: {jd_file}")
            except Exception as e:
                logging.error(f"处理职位说明书失败: {jd_file} - {str(e)}")
                continue

        return self.job_cache
    
class ResumeProcessor:
    """简历处理器"""
    def __init__(self):
        self.config = ConfigManager()
        self._setup_environment()

    def _setup_environment(self):
        tesseract_path = self.config.get('PATHS', 'tesseract_path')
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        poppler_path = self.config.get('PATHS', 'poppler_path')
        if os.path.exists(poppler_path):
            os.environ["PATH"] += os.pathsep + poppler_path

    def extract_text_from_pdf(self, pdf_path: str) -> Optional[str]:
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
            text = text.strip()
            if text:
                return text
        except Exception as e:
            logging.warning(f"PyPDF2提取失败: {pdf_path} - {str(e)}")

        try:
            images = convert_from_path(pdf_path)
            ocr_text = ""
            for image in images:
                image = image.convert('L')
                image = image.point(lambda x: 0 if x < 140 else 255)
                ocr_text += pytesseract.image_to_string(image, lang='chi_sim+eng')
            return ocr_text.strip() or text
        except Exception as e:
            logging.error(f"OCR提取失败: {pdf_path} - {str(e)}")
            return text

    def extract_text_from_docx(self, docx_path: str) -> str:
        try:
            doc = Document(docx_path)
            text = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            return '\n'.join(text)
        except Exception as e:
            logging.error(f"DOCX提取失败: {docx_path} - {str(e)}")
            return ""

    def extract_resume_text(self, file_path: str) -> str:
        if file_path.endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        elif file_path.endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        return ""


class DeepSeekEvaluator:
    """评估器：负责简历信息提取和候选人评估"""
    
    def __init__(self):
        self.config = ConfigManager()
        self.client = OpenAI(
            api_key=self.config.get('API', 'api_key'),
            base_url=self.config.get('API', 'base_url')
        )
        self.retry_count = 3
        self.retry_delay = 2
    
    def _build_extraction_prompt(self, resume_text: str, filename: str) -> str:
        """构建简历信息提取的Prompt，age 返回字符串"""
        return f"""分析简历，提取以下信息：

基本信息：
- 姓名：中文，从简历标题、个人信息、文件名提取。
- 性别：男/女，未明确返回""。
- 应聘职位：从简历标题、求职意向、文件名提取，未找到返回""。
- 年龄：从出生年份或年龄提取（如“30岁”），返回字符串，未找到返回""。
- 居住地：从联系方式、地址提取，未找到返回""。

教育背景：
- 原始学历：最早学历（学位、学校、专业、毕业年份），未找到返回空对象。
- 最高学历：最后学历（学位、学校、专业、毕业年份），未找到返回空对象。

工作经历：
- 按时间倒序，每段包含：
  - 时间段：YYYY年MM月-YYYY年MM月，非标准（如“2018-2020”）转为标准，至今转为“YYYY年MM月-2025年04月”，缺失标记“未知”。
  - 公司：名称，移除部门（如“腾讯技术部” -> “腾讯”）。
  - 性质：国企/民营/外企/合资/事业单位/其他，未明确返回“其他”。  
  - 规模：少于50人/50-100人/100-500人/500-1000人/1000人以上/未知，未明确返回“未知”。
  - 行业：互联网/软件/制造业/金融服务/教育培训/医疗健康/快消品/零售/物流/房地产/耐用消费品/进出口贸易/电子商务/其他，未明确返回“其他”。
  - 职位：名称，移除修饰（如“高级商务经理” -> “商务经理”）。
  - 描述：职责，限60字，超长截断。
- 优先从“工作经历”“职业经历”段落或表格提取，排除项目或实习经历。
- 推断规则：
  - 性质：如“中国石油” -> 国企，“腾讯” -> 民营。
  - 规模：如“初创” -> 少于50人，“上市公司” -> 1000人以上。
  - 行业：如“腾讯” -> 互联网，“宝洁” -> 快消品，“顺丰” -> 物流，“万达” -> 房地产，“美的” -> 耐用消费品，“中远海运” -> 进出口贸易，“京东” -> 电子商务。
- 未找到返回空列表。

项目经验：
- 仅从标有“项目经历”或“项目经验”的段落或表格提取，按时间倒序，每段包含：
  - 时间段：YYYY年MM月-YYYY年MM月，非标准转为标准，缺失标记“未知”。
  - 项目名称：如“商务平台开发”。
  - 角色：如“商务经理”“项目负责人”。
  - 技术栈：技术或工具，如“Excel, CRM”，未明确返回“”。
  - 项目成果：成果或影响，限60字，如“提升20%客户转化率”，未明确返回“”。
  - 描述：职责，限60字，超长截断。
- 未找到返回空列表。

技能及优势：
- 硬技能、软技能、个人优势列表，如“商务谈判”“团队协作”“抗压能力”。
- 硬技能可选掌握程度（精通/熟练/了解），软技能/优势留空。
- 从“技能”“自我评价”“工作经历”段落或表格提取。
- 未找到返回空列表。

简历文件名：{filename}
简历内容：{resume_text}

返回JSON：
{{
    "name": "",
    "gender": "",
    "position": "",
    "age": "",
    "location": "",
    "education": {{
        "original": {{"degree": "", "school": "", "major": "", "graduation_year": ""}},
        "highest": {{"degree": "", "school": "", "major": "", "graduation_year": ""}}
    }},
    "experience": {{
        "work_history": [
            {{
                "period": "",
                "company": "",
                "company_nature": "",                
                "company_scale": "",
                "company_industry": "",
                "position": "",
                "description": ""
            }}
        ]
    }},
    "projects": {{
        "project_history": [
            {{
                "period": "",
                "project_name": "",
                "role": "",
                "tech_stack": "",
                "outcomes": "",
                "description": ""
            }}
        ]
    }},
    "skills_and_strengths": {{
        "list": [],
        "proficiency": {{}}
    }}
}}
"""

    def _extract_resume_info(self, resume_text: str, filename: str) -> Dict:
        """从简历中提取信息"""
        prompt = self._build_extraction_prompt(resume_text, filename)
        last_error = None
        
        for attempt in range(self.retry_count):
            try:
                if attempt > 0:
                    time.sleep(self.retry_delay * (attempt + 1))
                
                response = self.client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[
                        {
                            "role": "system",
                            "content": "你是一个专业的简历信息提取专家。请严格按照要求格式提取信息，保持客观准确。"
                        },
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1,
                    max_tokens=3000,
                    response_format={"type": "json_object"}
                )
                
                content = response.choices[0].message.content
                extracted_info = self._parse_api_response(content, filename)
                return self._ensure_required_fields(extracted_info, filename)
            except Exception as e:
                last_error = str(e)
                logging.warning(f"API调用失败 (尝试 {attempt+1}/{self.retry_count}): {str(e)}")
                if attempt < self.retry_count - 1:
                    continue
                logging.error(f"所有提取尝试均失败: {filename} - {last_error}")
                return {}
            
        return {}

    def _parse_api_response(self, content: str, filename: str) -> Dict:
        """解析API响应内容，确保 age 是字符串"""
        try:
            content = content.strip()
            if content.startswith("```json"):
                content = content[7:-3].strip()
            elif content.startswith("```"):
                content = content[3:-3].strip()
            
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            try:
                result = json.loads(content)
                if not isinstance(result, dict):
                    logging.error(f"API响应非字典: {filename} - 类型: {type(result)}, 内容: {content[:500]}...")
                    return {}
            except json.JSONDecodeError as e:
                logging.error(f"JSON解析失败 ({filename}): {str(e)}\n原始内容: {content[:500]}...")
                return {}
            
            # 验证基本信息
            name = result.get('name', '').strip()
            if not name or not all('\u4e00' <= char <= '\u9fff' for char in name):
                logging.warning(f"提取的姓名无效或非中文 ({filename}): {name}")
                result['name'] = ''
            
            gender = result.get('gender', '').strip()
            if gender not in ['男', '女', '']:
                logging.warning(f"提取的性别值无效 ({filename}): {gender}")
                result['gender'] = ''
            
            position = result.get('position', '').strip()
            if not position:
                logging.warning(f"未提取到应聘职位 ({filename})")
            
            # 验证年龄
            age = result.get('age', '')
            if isinstance(age, dict):
                logging.info(f"年龄字段为字典，提取 summary ({filename}): {age}")
                age = age.get('summary', '')
            if not isinstance(age, str):
                logging.warning(f"年龄字段格式错误 ({filename}): {age}")
                age = ''
            result['age'] = age
            
            # 验证教育背景
            education = result.get('education', {})
            if not isinstance(education, dict):
                logging.warning(f"教育背景格式错误 ({filename}): {education}")
                education = {
                    'original': {'degree': '', 'school': '', 'major': '', 'graduation_year': ''},
                    'highest': {'degree': '', 'school': '', 'major': '', 'graduation_year': ''}
                }
            for edu_type in ['original', 'highest']:
                if not isinstance(education.get(edu_type), dict):
                    logging.warning(f"教育背景子项格式错误 ({filename}, {edu_type}): {education.get(edu_type)}")
                    education[edu_type] = {'degree': '', 'school': '', 'major': '', 'graduation_year': ''}
            
            # 验证工作经历
            experience = result.get('experience', {})
            if not isinstance(experience, dict):
                logging.warning(f"工作经历格式错误 ({filename}): {experience}")
                experience = {'work_history': []}
            work_history = experience.get('work_history', [])
            if not isinstance(work_history, list):
                logging.warning(f"工作经历列表格式错误 ({filename}): {work_history}")
                work_history = []
                experience['work_history'] = work_history
            valid_industries = [
                "互联网", "软件", "制造业", "金融服务", "教育培训", "医疗健康",
                "快消品", "零售", "物流", "房地产", "耐用消费品", "进出口贸易",
                "电子商务", "其他"
            ]
            valid_work_history = []
            for work in work_history:
                if not isinstance(work, dict):
                    logging.warning(f"工作经历项格式错误 ({filename}): {work}")
                    continue
                work["company_nature"] = str(work.get("company_nature", "其他"))             
                work["company_scale"] = str(work.get("company_scale", "未知"))
                work["company_industry"] = str(work.get("company_industry", "其他"))
                if work["company_industry"] not in valid_industries:
                    work["company_industry"] = "其他"
                work["period"] = str(work.get("period", ""))
                work["company"] = str(work.get("company", ""))
                work["position"] = str(work.get("position", ""))
                work["description"] = str(work.get("description", ""))[:60]
                valid_work_history.append(work)
            experience['work_history'] = valid_work_history
            
            # 验证项目经验
            projects = result.get('projects', {})
            if not isinstance(projects, dict):
                logging.warning(f"项目经验格式错误 ({filename}): {projects}")
                projects = {'project_history': []}
            project_history = projects.get('project_history', [])
            if not isinstance(project_history, list):
                logging.warning(f"项目经验列表格式错误 ({filename}): {project_history}")
                project_history = []
                projects['project_history'] = project_history
            valid_project_history = []
            for project in project_history:
                if not isinstance(project, dict):
                    logging.warning(f"项目经验项格式错误 ({filename}): {project}")
                    continue
                project["period"] = str(project.get("period", ""))
                project["project_name"] = str(project.get("project_name", ""))
                project["role"] = str(project.get("role", ""))
                project["tech_stack"] = str(project.get("tech_stack", ""))
                project["outcomes"] = str(project.get("outcomes", ""))[:60]
                project["description"] = str(project.get("description", ""))[:60]
                valid_project_history.append(project)
            projects['project_history'] = valid_project_history
            
            # 验证技能及优势
            skills = result.get('skills_and_strengths', {"list": [], "proficiency": {}})
            if not isinstance(skills, dict):
                logging.warning(f"技能及优势格式错误 ({filename}): {skills}")
                skills = {"list": [], "proficiency": {}}
            skills_list = skills.get('list', [])
            if not isinstance(skills_list, list):
                logging.warning(f"技能列表格式错误 ({filename}): {skills_list}")
                skills['list'] = []
            valid_skills_list = [str(skill) for skill in skills_list if isinstance(skill, str)]
            skills['list'] = valid_skills_list
            proficiency = skills.get('proficiency', {})
            if not isinstance(proficiency, dict):
                logging.warning(f"技能熟练度格式错误 ({filename}): {proficiency}")
                proficiency = {}
            valid_proficiency = {str(k): str(v) for k, v in proficiency.items() if isinstance(k, str) and isinstance(v, str)}
            skills['proficiency'] = valid_proficiency
            
            result['education'] = education
            result['experience'] = experience
            result['projects'] = projects
            result['skills_and_strengths'] = skills
            
            return result
        except Exception as e:
            logging.error(f"解析API响应失败 ({filename}): {str(e)}\n原始内容: {content[:500]}...")
            return {}
    
    def _ensure_required_fields(self, info: Dict, filename: str) -> Dict:
        """确保必要字段存在，age 为字符串"""
        try:
            filename_no_ext = os.path.splitext(filename)[0]
            filename_parts = filename_no_ext.split('_')
        
            platform = filename_parts[0] if len(filename_parts) > 0 else ""
            name_from_file = filename_parts[1] if len(filename_parts) > 1 else ""
            position_from_file = filename_parts[2] if len(filename_parts) >= 3 else filename_no_ext
        
            name_from_content = info.get('name', '').strip()
            final_name = ''
            if name_from_content and all('\u4e00' <= char <= '\u9fff' for char in name_from_content):
                final_name = name_from_content
            elif name_from_file and all('\u4e00' <= char <= '\u9fff' for char in name_from_file):
                final_name = name_from_file
                logging.info(f"使用文件名中的姓名: {final_name} ({filename})")
        
            position_from_content = info.get('position', '').strip()
            final_position = position_from_content or position_from_file
            if not position_from_content and position_from_file:
                logging.info(f"使用文件名中的职位: {final_position} ({filename})")
        
            gender_from_content = info.get('gender', '').strip()
            final_gender = ''
            if gender_from_content in ['男', '女']:
                final_gender = gender_from_content
            elif final_name:
                final_gender = self._infer_gender_from_name(final_name)
                if final_gender:
                    logging.info(f"从姓名推断性别: {final_gender} ({filename})")
         
            return {
                'name': final_name,
                'position': final_position,
                'gender': final_gender,
                'location': info.get('location', ''),
                'age': info.get('age', ''),
                'education': info.get('education', {
                    'original': {'degree': '', 'school': '', 'major': '', 'graduation_year': ''},
                    'highest': {'degree': '', 'school': '', 'major': '', 'graduation_year': ''}
                }),
                'experience': info.get('experience', {
                    'work_history': []
                }),
                'projects': info.get('projects', {
                    'project_history': []
                }),
                'skills_and_strengths': info.get('skills_and_strengths', {
                    'list': [],
                    'proficiency': {}
                })
            }
        except Exception as e:
            logging.error(f"处理字段补充失败 ({filename}): {str(e)}")
            return info
        
    def _build_result_dict(self, info: Dict, conclusion: str, filename: str) -> Dict:
        """构建最终结果字典，age 为字符串"""
        try:
            # 验证 info 结构
            if not isinstance(info, dict):
                logging.error(f"info 不是字典 ({filename}): {info}")
                return {}

            # 技能及优势
            skills = info.get('skills_and_strengths', {'list': [], 'proficiency': {}})
            if not isinstance(skills, dict):
                logging.warning(f"技能及优势格式错误 ({filename}): {skills}")
                skills = {'list': [], 'proficiency': {}}
            skills_list = skills.get('list', [])
            if not isinstance(skills_list, list):
                logging.warning(f"技能列表格式错误 ({filename}): {skills_list}")
                skills_list = []
            proficiency = skills.get('proficiency', {})
            if not isinstance(proficiency, dict):
                logging.warning(f"技能熟练度格式错误 ({filename}): {proficiency}")
                proficiency = {}
            skills_str = "; ".join([
                f"{skill} ({proficiency.get(skill, '')})" if proficiency.get(skill) else skill
                for skill in skills_list if isinstance(skill, str)
            ]) if skills_list else ""

            # 工作经历
            experience = info.get('experience', {'work_history': []})
            if not isinstance(experience, dict):
                logging.warning(f"工作经历格式错误 ({filename}): {experience}")
                experience = {'work_history': []}
            work_history = experience.get('work_history', [])
            if not isinstance(work_history, list):
                logging.warning(f"工作经历列表格式错误 ({filename}): {work_history}")
                work_history = []
            valid_work_history = []
            for exp in work_history:
                if not isinstance(exp, dict):
                    logging.warning(f"工作经历项格式错误 ({filename}): {exp}")
                    continue
                valid_exp = {
                    'period': str(exp.get('period', '')),
                    'company': str(exp.get('company', '')),
                    'company_nature': str(exp.get('company_nature', '其他')),
                    'company_scale': str(exp.get('company_scale', '未知')),
                    'company_industry': str(exp.get('company_industry', '其他')),
                    'position': str(exp.get('position', '')),
                    'description': str(exp.get('description', ''))[:60]
                }
                valid_work_history.append(valid_exp)
            work_history_str = '\n '.join([
                f"{exp['period']} {exp['company']} ({exp['company_nature']}, "
                f"{exp['company_scale']}, {exp['company_industry']}) {exp['position']}: {exp['description']}"
                for exp in valid_work_history
            ]) if valid_work_history else ''

            # 项目经验
            projects = info.get('projects', {'project_history': []})
            if not isinstance(projects, dict):
                logging.warning(f"项目经验格式错误 ({filename}): {projects}")
                projects = {'project_history': []}
            project_history = projects.get('project_history', [])
            if not isinstance(project_history, list):
                logging.warning(f"项目经验列表格式错误 ({filename}): {project_history}")
                project_history = []
            valid_project_history = []
            for proj in project_history:
                if not isinstance(proj, dict):
                    logging.warning(f"项目经验项格式错误 ({filename}): {proj}")
                    continue
                valid_proj = {
                    'period': str(proj.get('period', '')),
                    'project_name': str(proj.get('project_name', '')),
                    'role': str(proj.get('role', '')),
                    'tech_stack': str(proj.get('tech_stack', '')),
                    'outcomes': str(proj.get('outcomes', ''))[:60],
                    'description': str(proj.get('description', ''))[:60]
                }
                valid_project_history.append(valid_proj)
            project_history_str = '; '.join([
                f"{proj['period']} {proj['project_name']} ({proj['role']}, 技术栈: {proj['tech_stack']}, "
                f"成果: {proj['outcomes']}): {proj['description']}"
                for proj in valid_project_history
            ]) if valid_project_history else ''

            # 教育背景
            education = info.get('education', {
                'original': {'degree': '', 'school': '', 'major': '', 'graduation_year': ''},
                'highest': {'degree': '', 'school': '', 'major': '', 'graduation_year': ''}
            })
            if not isinstance(education, dict):
                logging.warning(f"教育背景格式错误 ({filename}): {education}")
                education = {
                    'original': {'degree': '', 'school': '', 'major': '', 'graduation_year': ''},
                    'highest': {'degree': '', 'school': '', 'major': '', 'graduation_year': ''}
                }

            # 构建结果
            return {
                '姓名': str(info.get('name', '')),
                '应聘岗位': str(info.get('position', '')),
                '性别': str(info.get('gender', '')),
                '居住地': str(info.get('location', '')),
                '年龄': str(info.get('age', '')),
                '原生学历': f"{education.get('original', {}).get('degree', '')} - "
                           f"{education.get('original', {}).get('school', '')} - "
                           f"{education.get('original', {}).get('major', '')} "
                           f"({education.get('original', {}).get('graduation_year', '')})",
                '最高学历': f"{education.get('highest', {}).get('degree', '')} - "
                           f"{education.get('highest', {}).get('school', '')} - "
                           f"{education.get('highest', {}).get('major', '')} "
                           f"({education.get('highest', {}).get('graduation_year', '')})",
                '工作经历': work_history_str,
                '项目经验': project_history_str,
                '技能及优势': skills_str,
                '评估结论': str(conclusion),
                '处理时间': datetime.now().strftime('%Y-%m-%d %H:%M'),
                '文件名': str(filename)
            }
        except Exception as e:
            logging.error(f"构建结果字典失败 ({filename}): {str(e)}\ninfo结构: {str(info)[:1000]}...")
            return {}
    
    def _infer_gender_from_name(self, name: str) -> str:
        """从姓名推断性别"""
        male_indicators = ['先生', '男', 'mr', 'mr.', '小哥']
        female_indicators = ['女士', '小姐', '女', 'ms', 'ms.', 'mrs', 'mrs.', '阿姨', '姐姐', '妹妹']
    
        name_lower = name.lower().strip()
        for indicator in male_indicators:
            if indicator.lower() in name_lower:
                return '男'
        for indicator in female_indicators:
            if indicator.lower() in name_lower:
                return '女'
        return ''
    
    def evaluate_candidate(self, resume_info: Dict, job_content: str, filename: str) -> str:
        """基于职位说明书内容评估候选人"""
        skills = resume_info.get('skills_and_strengths', {'list': [], 'proficiency': {}})
        skills_str = "; ".join([
            f"{skill} ({skills['proficiency'].get(skill, '')})" if skills['proficiency'].get(skill) else skill
            for skill in skills['list']
        ]) if skills['list'] else ""
        
        prompt = """根据以下职位说明书和候选人简历信息，评估候选人是否适合该岗位。
要求：
1. 比较教育背景、工作经历、项目经验和技能与岗位要求。
2. 输出简洁的评估结论（不超过100字）。
3. 结论需明确指出匹配度及主要优劣势。

职位说明书：
{job_content}

候选人信息：
姓名：{name}
教育背景：{education}
工作经历：{experience}
项目经验：{projects}
技能及优势：{skills}

返回格式：
评估结论：[具体结论]

示例返回：
评估结论：候选人技能匹配度高，10年商务经验符合要求，但学历略低于预期。
""".format(
            job_content=job_content,
            name=resume_info.get('name', ''),
            education=f"{resume_info.get('education', {}).get('highest', {}).get('degree', '')} - "
                      f"{resume_info.get('education', {}).get('highest', {}).get('school', '')} - "
                      f"{resume_info.get('education', {}).get('highest', {}).get('major', '')}",
            experience="; ".join([
                f"{exp['period']} {exp['company']} ({exp['company_nature']}, {exp['company_industry']}) "
                f"{exp['position']}: {exp['description']}"
                for exp in resume_info.get('experience', {}).get('work_history', []) if isinstance(exp, dict)
            ]),
            projects="; ".join([
                f"{proj['period']} {proj['project_name']} ({proj['role']}): {proj['description']}"
                for proj in resume_info.get('projects', {}).get('project_history', []) if isinstance(proj, dict)
            ]),
            skills=skills_str
        )

        try:
            response = self.client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "你是一个专业的招聘评估专家。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=150
            )
            content = response.choices[0].message.content.strip()
            if content.startswith("评估结论："):
                conclusion = content.replace("评估结论：", "").strip()
                logging.info(f"评估结论生成: {conclusion[:50]}... ({filename})")
                return conclusion
            return "评估结论无效"
        except Exception as e:
            logging.error(f"候选人评估失败: {filename} - {str(e)}")
            return "评估失败"
    
    def process_resume(self, resume_text: str, filename: str, job_cache: Dict[str, Dict[str, str]]) -> Dict:
        """提取简历信息并评估候选人"""
        try:
            info = self._extract_resume_info(resume_text, filename)
            if not info or not info.get('name'):
                logging.warning(f"简历信息提取失败或姓名为空: {filename}")
                return {}

            matched_jd_file, matched_position = self._match_position(
                resume_position=info.get('position', ''),
                filename=filename,
                job_cache=job_cache
            )

            if matched_position:
                info['position'] = matched_position
                logging.info(f"职位匹配成功: {matched_position} ({filename})")
            else:
                logging.warning(f"未匹配到职位: {filename}")

            if matched_jd_file and matched_jd_file in job_cache:
                conclusion = self.evaluate_candidate(
                    resume_info=info,
                    job_content=job_cache[matched_jd_file]['content'],
                    filename=filename
                )
            else:
                conclusion = "未匹配到岗位，无法评估"

            return self._build_result_dict(info, conclusion, filename)

        except Exception as e:
            logging.error(f"处理简历失败: {filename} - {str(e)}")
            return {}
    
    def _match_position(self, resume_position: str, filename: str, job_cache: Dict[str, Dict[str, str]]) -> Tuple[str, str]:
        """匹配职位名称"""
        try:
            clean_positions = []
            jd_file_mapping = {}
            for jd_file, data in job_cache.items():
                clean_position = data['position']
                clean_positions.append(clean_position)
                jd_file_mapping[clean_position] = jd_file

            clean_filename = re.sub(r'^\d+_|_\d+$|_南宁\(\d+\)$|_桂林\(\d+\)$|_昆明.*$|_.*\(\d+\)$', '', filename)
            clean_filename = re.sub(r'^51job_|^BOSS_|^zhilian_|^【|】$|\d+年以上|\d+-\d+K', '', clean_filename).strip()
            logging.info(f"清理后的文件名: {clean_filename} ({filename})")

            filename_parts = clean_filename.split('_')
            position_from_filename = filename_parts[2] if len(filename_parts) >= 3 else clean_filename

            if resume_position.strip():
                best_match = None
                highest_similarity = 0.0
                for clean_position in clean_positions:
                    similarity = self._calculate_position_similarity(resume_position, clean_position)
                    if similarity > highest_similarity:
                        highest_similarity = similarity
                        best_match = clean_position
                if highest_similarity > 0.7:
                    logging.info(f"简历内容匹配: {resume_position} -> {best_match} (相似度: {highest_similarity})")
                    return jd_file_mapping[best_match], best_match

            for clean_position in clean_positions:
                if clean_position == position_from_filename or clean_position == clean_filename:
                    logging.info(f"文件名精确匹配: {clean_filename} -> {clean_position}")
                    return jd_file_mapping[clean_position], clean_position
                if clean_position in clean_filename and len(clean_position) > len(position_from_filename) * 0.8:
                    logging.info(f"文件名部分匹配: {clean_filename} -> {clean_position}")
                    return jd_file_mapping[clean_position], clean_position

            logging.warning(f"未找到匹配职位: {clean_filename} (原始文件名: {filename})")
            return "", ""
        except Exception as e:
            logging.error(f"职位匹配失败: {filename} - {str(e)}")
            return "", ""

    def _calculate_position_similarity(self, pos1: str, pos2: str) -> float:
        """计算岗位名称相似度，优先完整匹配"""
        pos1, pos2 = pos1.strip(), pos2.strip()
        if pos1 == pos2:
            return 1.0
        if pos1 in pos2 and len(pos1) >= len(pos2) * 0.8:
            return 0.9
        if pos2 in pos1 and len(pos2) >= len(pos1) * 0.8:
            return 0.9
        words1 = set(pos1.split())
        words2 = set(pos2.split())
        common_words = words1 & words2
        if common_words:
            return len(common_words) / max(len(words1), len(words2)) * 0.7
        return 0.0

class ExcelGenerator:
    """Excel生成器"""
    STANDARD_COLUMNS = [      
        ('姓名', '姓名'),
        ('应聘岗位', '应聘岗位'),
        ('性别', '性别'),
        ('居住地', '居住地'),
        ('年龄', '年龄'),
        ('原生学历', '原生学历'),
        ('最高学历', '最高学历'),
        ('工作经历', '工作经历'),
        ('项目经验', '项目经验'),
        ('技能及优势', '技能及优势'),
        ('评估结论', '评估结论'),
        ('处理时间', '处理时间'),
        ('文件名', '文件名')       
    ]
    
    @staticmethod
    def _validate_excel_headers(worksheet, expected_headers: List[str]) -> bool:
        """验证Excel表头是否与预期一致"""
        if not worksheet['A1'].value:  # 空表
            return True
        headers = [cell.value for cell in worksheet[1] if cell.value]
        return headers == expected_headers

    @staticmethod
    def generate(results: List[Dict], output_path: str) -> str:
        """追加结果到现有Excel文件或创建新文件"""
        try:
            # 验证output_path是否有效
            if not output_path.lower().endswith(('.xlsx', '.xlsm', '.xltx', '.xltm')):
                logging.error(f"无效的Excel文件路径: {output_path} - 必须是.xlsx, .xlsm, .xltx或.xltm格式")
                raise ValueError("输出Excel文件路径无效，必须是.xlsx格式")

            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            os.makedirs(output_dir, exist_ok=True)

            # 准备数据
            for result in results:
                if '评估结论' not in result or not result['评估结论']:
                    logging.warning(f"结果中缺失或空的评估结论: {result.get('文件名', '未知文件')}")

            df = pd.DataFrame(results)
            ordered_columns = [clean_col for _, clean_col in ExcelGenerator.STANDARD_COLUMNS]
            df = df.reindex(columns=ordered_columns)
            df.fillna('', inplace=True)

            # Excel格式设置
            min_width = 10
            content_alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
            header_fill = PatternFill(start_color='CCE5FF', end_color='CCE5FF', fill_type='solid')
            header_font = Font(bold=True)
            header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

            sheet_name = '简历信息'
            expected_headers = ordered_columns

            # 检查文件是否存在
            if os.path.exists(output_path):
                try:
                    # 加载现有Excel
                    workbook = openpyxl.load_workbook(output_path)
                    if sheet_name in workbook.sheetnames:
                        worksheet = workbook[sheet_name]
                        # 验证表头
                        if not ExcelGenerator._validate_excel_headers(worksheet, expected_headers):
                            logging.error(f"Excel文件表头不匹配: {output_path}")
                            raise ValueError(f"Excel文件 '{output_path}' 的表头与预期不匹配")
                    else:
                        # 创建新工作表
                        worksheet = workbook.create_sheet(sheet_name)
                except Exception as e:
                    logging.error(f"加载Excel文件失败: {output_path} - {str(e)}")
                    raise
            else:
                # 创建新Excel
                workbook = openpyxl.Workbook()
                worksheet = workbook.active
                worksheet.title = sheet_name

            # 写入表头（仅在新文件或新工作表时）
            if worksheet.max_row == 1 and not worksheet['A1'].value:
                for col_idx, header in enumerate(expected_headers, start=1):
                    cell = worksheet.cell(row=1, column=col_idx)
                    cell.value = header
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = header_alignment

            # 追加数据
            start_row = worksheet.max_row + 1 if worksheet.max_row > 1 else 2
            for row_idx, row_data in enumerate(df.values, start=start_row):
                for col_idx, value in enumerate(row_data, start=1):
                    cell = worksheet.cell(row=row_idx, column=col_idx)
                    cell.value = str(value)
                    cell.alignment = content_alignment

            # 调整列宽
            for col_idx, column in enumerate(df.columns, start=1):
                max_length = max(
                    df[column].astype(str).str.len().max(),
                    len(str(column))
                )
                adjusted_width = max(min_width, min(max_length * 2.5, 100))
                col_letter = get_column_letter(col_idx)
                worksheet.column_dimensions[col_letter].width = adjusted_width

            # 设置冻结窗格
            worksheet.freeze_panes = 'B2'

            # 保存文件
            workbook.save(output_path)
            logging.info(f"成功追加数据到Excel: {output_path} (新增 {len(results)} 行)")
            return output_path

        except Exception as e:
            logging.error(f"生成或追加Excel失败: {str(e)}", exc_info=True)
            raise

class RecruitmentSystemGUI:
    """招聘系统GUI界面"""
    def __init__(self, root):
        self.root = root
        self.root.title("招聘管理系统")
        self.root.geometry("1000x800")
        
        self.resume_processor = ResumeProcessor()
        self.evaluator = DeepSeekEvaluator()
        self.job_desc_processor = JobDescriptionProcessor()
        self.config = ConfigManager()
        
        self.main_frame = ttk.Frame(root, padding="10")
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        
        self.work_dir = tk.StringVar(value=self.config.get('PATHS', 'work_dir'))
        self.resume_dir = tk.StringVar(value=self.config.get('PATHS', 'resume_dir'))
        self.job_desc_dir = tk.StringVar(value=self.config.get('PATHS', 'job_desc_dir'))
        # 确保output_excel有默认有效路径
        work_dir = self.work_dir.get() or os.path.expanduser("~")
        default_excel = self.config.get('PATHS', 'output_excel')
        if not default_excel.lower().endswith(('.xlsx', '.xlsm', '.xltx', '.xltm')):
            default_excel = os.path.join(work_dir, '简历信息一览表.xlsx')
        self.output_excel = tk.StringVar(value=self.config.get('PATHS', 'output_excel') or default_excel)
        self.job_desc_files = []
        self.progress_var = tk.DoubleVar()
        self.status_var = tk.StringVar(value="准备就绪")
        self.running = False
        
        self._setup_ui()
        self._redirect_logging()

    def _setup_ui(self):
        """初始化用户界面"""
        dir_frame = ttk.LabelFrame(self.main_frame, text="目录设置", padding="10")
        dir_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 工作目录
        ttk.Label(dir_frame, text="工作目录:").grid(row=0, column=0, sticky=tk.W)
        ttk.Entry(dir_frame, textvariable=self.work_dir, width=70).grid(row=0, column=1, padx=5)
        ttk.Button(dir_frame, text="浏览", command=self.browse_work_dir).grid(row=0, column=2)
        
        # 简历目录
        ttk.Label(dir_frame, text="简历目录:").grid(row=1, column=0, sticky=tk.W)
        ttk.Entry(dir_frame, textvariable=self.resume_dir, width=70).grid(row=1, column=1, padx=5)
        ttk.Button(dir_frame, text="浏览", command=self.browse_resume_dir).grid(row=1, column=2)
        
        # 职位说明书目录
        ttk.Label(dir_frame, text="职位说明书目录:").grid(row=2, column=0, sticky=tk.W)
        ttk.Entry(dir_frame, textvariable=self.job_desc_dir, width=70).grid(row=2, column=1, padx=5)
        ttk.Button(dir_frame, text="浏览", command=self.browse_job_desc_dir).grid(row=2, column=2)
        
        # 输出Excel文件
        ttk.Label(dir_frame, text="输出Excel文件:").grid(row=3, column=0, sticky=tk.W)
        ttk.Entry(dir_frame, textvariable=self.output_excel, width=70).grid(row=3, column=1, padx=5)
        ttk.Button(dir_frame, text="浏览", command=self.browse_output_excel).grid(row=3, column=2)
        
        job_desc_frame = ttk.LabelFrame(self.main_frame, text="职位说明书", padding="10")
        job_desc_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(job_desc_frame, text="选择职位说明书:").grid(row=0, column=0, sticky=tk.W)
        
        self.job_desc_listbox = tk.Listbox(job_desc_frame, width=70, height=5, selectmode=tk.MULTIPLE)
        self.job_desc_listbox.grid(row=1, column=0, columnspan=2, padx=5, pady=5)
        
        button_frame = ttk.Frame(job_desc_frame)
        button_frame.grid(row=1, column=2, sticky=tk.N)
        
        ttk.Button(button_frame, text="添加", command=self.add_job_description).pack(fill=tk.X, pady=2)
        ttk.Button(button_frame, text="移除", command=self.remove_job_description).pack(fill=tk.X, pady=2)
        
        progress_frame = ttk.Frame(self.main_frame)
        progress_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(progress_frame, textvariable=self.status_var).pack(anchor=tk.W)
        ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100).pack(fill=tk.X)
        
        log_frame = ttk.LabelFrame(self.main_frame, text="系统日志", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True)
        
        self.log_text = tk.Text(log_frame, height=20, state='disabled')
        scrollbar = ttk.Scrollbar(log_frame, command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        button_frame = ttk.Frame(self.main_frame)
        button_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(
            button_frame, 
            text="处理简历并追加到Excel", 
            command=self.process_resumes_and_generate_excel,
            style='Accent.TButton'
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            button_frame, 
            text="查看日志文件", 
            command=self.view_logs
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            button_frame, 
            text="退出系统", 
            command=self.root.quit,
            style='Warning.TButton'
        ).pack(side=tk.RIGHT)

    def _redirect_logging(self):
        """重定向日志到GUI文本框"""
        class TextHandler(logging.Handler):
            def __init__(self, text_widget):
                super().__init__()
                self.text_widget = text_widget
            
            def emit(self, record):
                msg = self.format(record)
                self.text_widget.configure(state='normal')
                self.text_widget.insert(tk.END, msg + '\n')
                self.text_widget.configure(state='disabled')
                self.text_widget.see(tk.END)
        
        text_handler = TextHandler(self.log_text)
        text_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
        logging.getLogger().addHandler(text_handler)

    def browse_work_dir(self):
        """浏览选择工作目录"""
        directory = filedialog.askdirectory(initialdir=self.work_dir.get() or os.path.expanduser("~"))
        if directory:
            self.work_dir.set(directory)
            self.config.set('PATHS', 'work_dir', directory)
            # 更新默认Excel路径
            default_excel = os.path.join(directory, '简历信息一览表.xlsx')
            self.output_excel.set(default_excel)
            self.config.set('PATHS', 'output_excel', default_excel)
            self.status_var.set(f"已选择工作目录: {directory}")
            logging.info(f"工作目录设置为: {directory}")

    def browse_resume_dir(self):
        """浏览选择简历目录"""
        directory = filedialog.askdirectory(initialdir=self.resume_dir.get() or self.work_dir.get() or os.path.expanduser("~"))
        if directory:
            self.resume_dir.set(directory)
            self.config.set('PATHS', 'resume_dir', directory)
            self.status_var.set(f"已选择简历目录: {directory}")
            logging.info(f"简历目录设置为: {directory}")

    def browse_job_desc_dir(self):
        """浏览选择职位说明书目录"""
        directory = filedialog.askdirectory(initialdir=self.job_desc_dir.get() or self.work_dir.get() or os.path.expanduser("~"))
        if directory:
            self.job_desc_dir.set(directory)
            self.config.set('PATHS', 'job_desc_dir', directory)
            self.status_var.set(f"已选择职位说明书目录: {directory}")
            logging.info(f"职位说明书目录设置为: {directory}")

    def browse_output_excel(self):
        """浏览选择输出Excel文件"""
        file_path = filedialog.asksaveasfilename(
            initialdir=os.path.dirname(self.output_excel.get()) or self.work_dir.get() or os.path.expanduser("~"),
            defaultextension=".xlsx",
            filetypes=[("Excel文件", "*.xlsx"), ("所有文件", "*.*")],
            initialfile=os.path.basename(self.output_excel.get()) or "简历信息一览表.xlsx"
        )
        if file_path:
            if not file_path.lower().endswith(('.xlsx', '.xlsm', '.xltx', '.xltm')):
                messagebox.showerror("错误", "请选择有效的Excel文件（.xlsx格式）")
                return
            self.output_excel.set(file_path)
            self.config.set('PATHS', 'output_excel', file_path)
            self.status_var.set(f"已选择输出Excel文件: {file_path}")
            logging.info(f"输出Excel文件设置为: {file_path}")

    def add_job_description(self):
        """添加职位说明书"""
        job_desc_dir = self.job_desc_dir.get()
        if not job_desc_dir or not os.path.exists(job_desc_dir):
            messagebox.showwarning("警告", "请先选择职位说明书目录")
            return
        
        files = filedialog.askopenfilenames(
            initialdir=job_desc_dir,
            filetypes=[("Word文件", "*.docx"), ("所有文件", "*.*")]
        )
        for file in files:
            if file not in self.job_desc_files:
                self.job_desc_files.append(file)
                self.job_desc_listbox.insert(tk.END, os.path.basename(file))
        self.status_var.set(f"已添加 {len(files)} 个职位说明书")
        logging.info(f"添加职位说明书: {files}")

    def remove_job_description(self):
        """移除选中的职位说明书"""
        selections = self.job_desc_listbox.curselection()
        if selections:
            for index in reversed(selections):
                removed_file = self.job_desc_files.pop(index)
                self.job_desc_listbox.delete(index)
                logging.info(f"移除职位说明书: {removed_file}")
            self.status_var.set(f"已移除 {len(selections)} 个职位说明书")

    def view_logs(self):
        """查看日志文件"""
        log_file = 'recruitment_system.log'
        if os.path.exists(log_file):
            try:
                if os.name == 'nt':
                    os.startfile(log_file)
                else:
                    subprocess.run(['open', log_file])
            except Exception as e:
                messagebox.showerror("错误", f"无法打开日志文件: {str(e)}")
        else:
            messagebox.showinfo("提示", "日志文件不存在")

    def process_resumes_and_generate_excel(self):
        """处理简历、评估候选人并追加到Excel"""
        if self.running:
            return
    
        self.running = True
        work_dir = self.work_dir.get()
        resume_dir = self.resume_dir.get()
        job_desc_dir = self.job_desc_dir.get()
        output_excel = self.output_excel.get()
        
        if not work_dir:
            messagebox.showerror("错误", "请先选择工作目录")
            self.running = False
            return
        if not resume_dir:
            messagebox.showerror("错误", "请先选择简历目录")
            self.running = False
            return
        if not job_desc_dir:
            messagebox.showerror("错误", "请先选择职位说明书目录")
            self.running = False
            return
        if not output_excel or not output_excel.lower().endswith(('.xlsx', '.xlsm', '.xltx', '.xltm')):
            messagebox.showerror("错误", "请先选择有效的Excel文件（.xlsx格式）")
            self.running = False
            return
    
        try:
            if not os.path.exists(resume_dir):
                messagebox.showerror("错误", "简历目录不存在")
                self.running = False
                return
        
            resume_files = [f for f in os.listdir(resume_dir) if f.endswith(('.pdf', '.docx'))]
            if not resume_files:
                messagebox.showerror("错误", "简历目录中没有简历文件")
                self.running = False
                return
        
            processed_dir = os.path.join(work_dir, '已处理简历')
            os.makedirs(processed_dir, exist_ok=True)
        
            job_cache = self.job_desc_processor.process_job_descriptions(self.job_desc_files)
            if not job_cache:
                messagebox.showerror("错误", "未成功处理任何职位说明书")
                self.running = False
                return
        
            results = []
            total_files = len(resume_files)
            for i, filename in enumerate(resume_files):
                try:
                    self.progress_var.set((i + 1) / total_files * 100)
                    self.status_var.set(f"正在处理 {filename} ({i+1}/{total_files})")
                    self.root.update()
                
                    file_path = os.path.join(resume_dir, filename)
                    resume_text = self.resume_processor.extract_resume_text(file_path)
                
                    if not resume_text.strip():
                        logging.warning(f"简历内容为空: {filename}")
                        continue
                
                    info = self.evaluator.process_resume(
                        resume_text=resume_text,
                        filename=filename,
                        job_cache=job_cache
                    )
                
                    if info:
                        results.append(info)
                        logging.info(f"成功处理简历: {filename} - 评估结论: {info['评估结论'][:50]}...")
                
                    try:
                        dest_path = os.path.join(processed_dir, filename)
                        shutil.move(file_path, dest_path)
                        logging.info(f"已移动简历文件到: {dest_path}")
                    except Exception as e:
                        logging.error(f"移动简历文件失败: {filename} - {str(e)}")
                except Exception as e:
                    logging.error(f"处理简历失败: {filename} - {str(e)}")
                    continue
        
            if results:
                try:
                    output_path = ExcelGenerator.generate(results, output_excel)
                
                    messagebox.showinfo(
                        "处理完成",
                        f"成功处理 {len(results)}/{total_files} 份简历\n"
                        f"数据已追加到:\n{output_path}"
                    )
                
                    try:
                        if os.name == 'nt':
                            os.startfile(output_path)
                        else:
                            subprocess.run(['open', output_path])
                    except Exception as e:
                        logging.warning(f"无法打开结果文件: {str(e)}")
                        messagebox.showwarning("警告", f"无法自动打开结果文件，请手动打开：\n{output_path}")
                except ValueError as e:
                    messagebox.showerror("错误", str(e))
                    self.running = False
                    return
                except Exception as e:
                    messagebox.showerror("错误", f"无法保存Excel文件：{str(e)}\n请确保文件未被占用且路径有效")
                    self.running = False
                    return
            else:
                messagebox.showerror("错误", "没有成功处理任何简历")
    
        except Exception as e:
            messagebox.showerror("系统错误", f"处理过程中发生错误:\n{str(e)}")
            logging.error(f"系统错误: {str(e)}", exc_info=True)
        finally:
            self.running = False
            self.progress_var.set(100)
            self.status_var.set("处理完成")

def main():
    root = tk.Tk()
    try:
        root.iconbitmap('icon.ico')
    except:
        pass
    
    style = ttk.Style()
    style.theme_use('clam')
    style.configure('Accent.TButton', foreground='white', background='#0078d7')
    style.configure('Warning.TButton', foreground='white', background='#d73a49')
    
    app = RecruitmentSystemGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()